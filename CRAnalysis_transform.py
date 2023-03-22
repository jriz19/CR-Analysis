import pandas as pd
import requests
import re
from io import BytesIO
from docx import Document

# Download the data from the link
#url = "https://www.tse.go.cr/pdf/padron/sumaria_pcd.docx"
#response = requests.get(url)
#document = Document(BytesIO(response.content))

# Read the file from a local path
file_path = "C:\\Users\\josea\\Downloads\\sumaria_pcd.docx"
document = Document(file_path)

# Read the content and process the data
data = []
pattern = re.compile(r"(\d+)\s+(.+?)\s+([\d,]+)\s+([\d,]+)\s+([\d,]+)\s+(-?\d+)\s+([\d,]+)\s+(-?\d+)\s+(-?\d+\.\d+)\s+%")

for paragraph in document.paragraphs:
    line = paragraph.text.strip()
    match = pattern.match(line)
    
    if match:
        groups = match.groups()
        codigo = groups[0]
        district = groups[1]
        total = groups[2]
        hombres = groups[3]
        mujeres = groups [4]
        variacion_mes_anterior = groups [5]
        padron_eleccion_2022 = groups [6]
        variacion_absoluto = groups [7]
        variacion_relativo = groups [8]

        data.append([
            codigo,
            district,
            int(total),
            int(hombres),
            int(mujeres),
            int(variacion_mes_anterior),
            int(padron_eleccion_2022),
            int(variacion_absoluto),
            variacion_relativo
        ])

# Create a pandas DataFrame
columns = [
    "Codigo",
    "Distrito",
    "Total",
    "Hombres",
    "Mujeres",
    "Variacion_Mes_Anterior",
    "Padron_Eleccion_2022",
    "Variacion_Absoluto",
    "Variacion_Relativo"
]
df = pd.DataFrame(data, columns=columns)

# Save the DataFrame to a CSV file
df.to_csv("electoral_data.csv", index=False)

print("Data successfully saved to 'electoral_data.csv'")
